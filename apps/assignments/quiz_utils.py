import csv
import io
import pandas as pd
from docx import Document
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from .models import Assignments, QuizQuestions, QuizChoices
from core.decorators import teacher_required

@login_required
@teacher_required
def download_quiz_template_view(request, format):
    """Xuất file mẫu trắc nghiệm (csv, xlsx, docx)."""
    sample_data = [
        {
            'question': 'Kiểu dữ liệu nào trong Python là immutable?',
            'type': 'single_choice',
            'points': 1.0,
            'choice_1': 'list',
            'choice_2': 'dict',
            'choice_3': 'tuple',
            'choice_4': 'set',
            'correct_index': 3,
            'explanation': 'Tuple không thể thay đổi sau khi tạo.'
        }
    ]

    if format == 'csv':
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=sample_data[0].keys())
        writer.writeheader()
        writer.writerows(sample_data)
        response = HttpResponse(output.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="quiz_template.csv"'
        return response

    elif format == 'xlsx':
        df = pd.DataFrame(sample_data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="quiz_template.xlsx"'
        return response

    elif format == 'docx':
        doc = Document()
        doc.add_heading('Mẫu Import Câu hỏi Trắc nghiệm', 0)
        doc.add_paragraph('Vui lòng nhập theo cấu trúc bên dưới. Mỗi câu hỏi bắt đầu bằng [QUESTION].')
        
        for item in sample_data:
            doc.add_paragraph(f"[QUESTION]: {item['question']}")
            doc.add_paragraph(f"[TYPE]: {item['type']}")
            doc.add_paragraph(f"[POINTS]: {item['points']}")
            doc.add_paragraph(f"[A]: {item['choice_1']}")
            doc.add_paragraph(f"[B]: {item['choice_2']}")
            doc.add_paragraph(f"[C]: {item['choice_3']}")
            doc.add_paragraph(f"[D]: {item['choice_4']}")
            doc.add_paragraph(f"[CORRECT]: C")
            doc.add_paragraph(f"[EXPLANATION]: {item['explanation']}")
            doc.add_paragraph("-" * 20)

        output = io.BytesIO()
        doc.save(output)
        response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="quiz_template.docx"'
        return response

    return HttpResponse("Định dạng không hỗ trợ", status=400)

def parse_quiz_file(file_obj):
    """Parse CSV, XLSX, or DOCX into a list of standardized rows."""
    filename = file_obj.name.lower()
    rows = []
    
    if filename.endswith('.csv'):
        stream = io.StringIO(file_obj.read().decode('utf-8-sig'))
        reader = csv.DictReader(stream)
        for i, row in enumerate(reader, start=1):
            rows.append({
                'line_no': i,
                'question_text': row.get('question_text') or row.get('question'),
                'question_type': row.get('question_type') or row.get('type') or 'single_choice',
                'points': float(row.get('points') or 1.0),
                'choice_1': row.get('choice_a') or row.get('choice_1') or row.get('option_1'),
                'choice_2': row.get('choice_b') or row.get('choice_2') or row.get('option_2'),
                'choice_3': row.get('choice_c') or row.get('choice_3') or row.get('option_3'),
                'choice_4': row.get('choice_d') or row.get('choice_4') or row.get('option_4'),
                'correct_index': row.get('correct_answers') or row.get('correct_index'),
                'explanation': row.get('explanation', '')
            })
            
    elif filename.endswith('.xlsx'):
        df = pd.read_excel(file_obj)
        for i, row_data in df.iterrows():
            row = row_data.to_dict()
            rows.append({
                'line_no': i + 1,
                'question_text': row.get('question_text') or row.get('question'),
                'question_type': row.get('question_type') or row.get('type') or 'single_choice',
                'points': float(row.get('points') or 1.0),
                'choice_1': row.get('choice_a') or row.get('choice_1') or row.get('option_1'),
                'choice_2': row.get('choice_b') or row.get('choice_2') or row.get('option_2'),
                'choice_3': row.get('choice_c') or row.get('choice_3') or row.get('option_3'),
                'choice_4': row.get('choice_d') or row.get('choice_4') or row.get('option_4'),
                'correct_index': row.get('correct_answers') or row.get('correct_index'),
                'explanation': row.get('explanation', '')
            })
            
    elif filename.endswith('.docx'):
        doc = Document(file_obj)
        current_q = {}
        idx = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            
            if text.startswith('[QUESTION]:'):
                if current_q: rows.append(current_q)
                idx += 1
                current_q = {
                    'line_no': idx, 
                    'question_text': text.replace('[QUESTION]:', '').strip(),
                    'question_type': 'single_choice',
                    'points': 1.0,
                    'choices': [],
                    'correct_answers': [],
                    'explanation': '',
                }
            elif not current_q: continue
            elif text.startswith('[TYPE]:'): current_q['question_type'] = text.replace('[TYPE]:', '').strip()
            elif text.startswith('[A]:'): current_q['choice_1'] = text.replace('[A]:', '').strip()
            elif text.startswith('[B]:'): current_q['choice_2'] = text.replace('[B]:', '').strip()
            elif text.startswith('[C]:'): current_q['choice_3'] = text.replace('[C]:', '').strip()
            elif text.startswith('[D]:'): current_q['choice_4'] = text.replace('[D]:', '').strip()
            elif text.startswith('[CORRECT]:'):
                val = text.replace('[CORRECT]:', '').strip().upper()
                mapping = {'A': 1, 'B': 2, 'C': 3, 'D': 4, '1': 1, '2': 2, '3': 3, '4': 4}
                current_q['correct_index'] = mapping.get(val, 1)
            elif text.startswith('[EXPLANATION]:'): current_q['explanation'] = text.replace('[EXPLANATION]:', '').strip()
        if current_q: rows.append(current_q)
        
    return rows

def process_quiz_import(file_obj, assignment):
    """Xử lý parse file và lưu vào database."""
    try:
        rows = parse_quiz_file(file_obj)
        total_questions = len(rows)
        if total_questions == 0:
            return False, "File không chứa câu hỏi hợp lệ."
        
        points_per_question = round(100.0 / total_questions, 2)
        
        for row in rows:
            row['points'] = points_per_question
            _create_question_from_row(row, assignment)
            
        return True, total_questions
    except Exception as e:
        return False, str(e)

def _create_question_from_row(row, assignment):
    """Helper để lưu câu hỏi và đáp án."""
    q = QuizQuestions.objects.create(
        assignment=assignment,
        question_text=row.get('question') or row.get('question_text'),
        question_type=row.get('type') or 'single_choice',
        points=float(row.get('points') or 1.0),
        explanation=row.get('explanation', ''),
        is_active=True
    )
    
    # Tạo đáp án (giả định mẫu có 4 lựa chọn)
    for i in range(1, 5):
        choice_text = row.get(f'choice_{i}') or row.get(f'option_{i}')
        if choice_text:
            QuizChoices.objects.create(
                question=q,
                choice_text=choice_text,
                is_correct=(int(row.get('correct_index') or 0) == i),
                order_index=i
            )
