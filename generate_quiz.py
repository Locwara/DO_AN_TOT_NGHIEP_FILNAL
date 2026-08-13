import docx

questions = [
    {
        "q": "Ngôn ngữ lập trình Python do ai tạo ra?",
        "a": "Guido van Rossum",
        "b": "James Gosling",
        "c": "Dennis Ritchie",
        "d": "Bjarne Stroustrup",
        "explain": "Python được Guido van Rossum tạo ra và phát hành lần đầu vào năm 1991."
    },
    {
        "q": "Đâu là cách đúng để tạo một chú thích (comment) một dòng trong Python?",
        "a": "# Đây là chú thích",
        "b": "// Đây là chú thích",
        "c": "/* Đây là chú thích */",
        "d": "<!-- Đây là chú thích -->",
        "explain": "Python sử dụng ký tự # cho các chú thích trên một dòng."
    },
    {
        "q": "Để in ra màn hình chuỗi 'Hello World' trong Python, câu lệnh nào sau đây là đúng?",
        "a": "print('Hello World')",
        "b": "echo('Hello World')",
        "c": "printf('Hello World')",
        "d": "cout << 'Hello World'",
        "explain": "Hàm print() được dùng để xuất dữ liệu ra màn hình trong Python."
    },
    {
        "q": "Kiểu dữ liệu nào sau đây được dùng để lưu trữ danh sách các giá trị có thể thay đổi?",
        "a": "List",
        "b": "Tuple",
        "c": "String",
        "d": "Integer",
        "explain": "List là một tập hợp các phần tử có thứ tự và có thể thay đổi (mutable)."
    },
    {
        "q": "Từ khóa nào được sử dụng để định nghĩa một hàm trong Python?",
        "a": "def",
        "b": "function",
        "c": "fun",
        "d": "define",
        "explain": "Trong Python, các hàm được khai báo bằng từ khóa def."
    },
    {
        "q": "Toán tử nào được dùng để tính lũy thừa trong Python?",
        "a": "**",
        "b": "^",
        "c": "//",
        "d": "%%",
        "explain": "Toán tử ** dùng để tính lũy thừa. Ví dụ 2**3 = 8."
    },
    {
        "q": "Hàm nào sau đây trả về độ dài của một danh sách (list)?",
        "a": "len()",
        "b": "length()",
        "c": "size()",
        "d": "count()",
        "explain": "Hàm len() trả về số lượng phần tử của một đối tượng có thể lặp (iterable)."
    },
    {
        "q": "Đoạn mã `type(10.5)` sẽ trả về kiểu dữ liệu gì?",
        "a": "float",
        "b": "int",
        "c": "double",
        "d": "str",
        "explain": "10.5 là một số thực, do đó type(10.5) trả về float."
    },
    {
        "q": "Câu lệnh nào dùng để nhập dữ liệu từ bàn phím?",
        "a": "input()",
        "b": "scanf()",
        "c": "read()",
        "d": "get()",
        "explain": "Hàm input() cho phép đọc một dòng văn bản nhập vào từ người dùng."
    },
    {
        "q": "Vòng lặp nào sau đây duyệt qua các phần tử của một danh sách?",
        "a": "for item in my_list:",
        "b": "for (i=0; i<my_list.length; i++)",
        "c": "foreach item in my_list:",
        "d": "loop item in my_list:",
        "explain": "Trong Python, vòng lặp for in được dùng để duyệt qua các iterable như list, tuple."
    }
]

doc = docx.Document()
for i, q in enumerate(questions):
    doc.add_paragraph(f"[QUESTION]: Câu {i+1}: {q['q']}")
    doc.add_paragraph("[TYPE]: single_choice")
    doc.add_paragraph(f"[A]: {q['a']}")
    doc.add_paragraph(f"[B]: {q['b']}")
    doc.add_paragraph(f"[C]: {q['c']}")
    doc.add_paragraph(f"[D]: {q['d']}")
    doc.add_paragraph("[CORRECT]: A")
    doc.add_paragraph(f"[EXPLANATION]: {q['explain']}")
    doc.add_paragraph("")

file_path = "/home/locwara/DO_AN_TOT_NGHIEP_FINAL/src/Websitedayvahoclaptrinh/quyen_bao_cao/quiz_python_10cau.docx"
doc.save(file_path)
print("Docx created:", file_path)
